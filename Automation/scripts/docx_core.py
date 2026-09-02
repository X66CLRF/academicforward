# -*- coding: utf-8 -*-
"""docx_core — ไลบรารีกลางสำหรับสร้างไฟล์ Word ภาษาไทย (และภาษาอักษรเชิงซ้อนอื่น)

ใช้คู่กับคู่มือ Automation/docx-builder.md

ทำไมต้องมีไฟล์นี้
    แม่แบบเปล่าของ python-docx มีปัญหากับภาษาไทย 3 จุด ซึ่งตรวจด้วยตาไม่เจอ
    ต้องเปิด XML ดูถึงจะรู้ ไฟล์นี้แก้ไว้ให้ครบแล้ว
      1. w:bidi ค่าเริ่มต้นเป็น ar-SA (อาหรับ) ทำให้ตรวจคำสะกดผิดภาษา
      2. ไม่มีธง <w:cs/> ระดับ run ทำให้ Word ตัดบรรทัดไทยผิดตำแหน่ง
         (Word ไม่สืบทอดธงนี้จากสไตล์ ต้องเขียนลงทุก run)
      3. ไม่มี <w:applyBreakingRules/> ซึ่งเป็นสวิตช์เปิดกฎการตัดบรรทัด
    บวกกับหลักการ "ผูกทุกค่าไว้ที่ Word Styles ห้ามจัดรูปแบบโดยตรง"
    เพื่อให้ผู้ใช้แก้สไตล์ตัวเดียวแล้วเปลี่ยนทั้งเล่ม

ไม่ผูกกับสถาบัน ประเภทงาน หรือภาษาใด ค่าทุกตัวส่งเข้ามาทาง SPEC ได้

การใช้งานอย่างย่อ
    from docx_core import DocBuilder, SPEC_TH_ACADEMIC

    b = DocBuilder(SPEC_TH_ACADEMIC)
    b.p("บทที่ 1", "หัวบท")
    b.p("ชื่อบท", "ชื่อบท")
    b.body("ย่อหน้าเนื้อความ...")
    b.p("1. ข้อระดับหนึ่ง", "ข้อ1")
    b.table("ตารางที่ 1 ชื่อตาราง", ["คอลัมน์ ก", "คอลัมน์ ข"], [(1, 2), (3, 4)])
    b.save("output.docx")

ต้องมี  pip install python-docx
"""

from __future__ import annotations

import os
import copy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

__all__ = ["DocBuilder", "SPEC_TH_ACADEMIC", "verify", "measure"]


# ── ลำดับลูกของ w:rPr ตามที่ schema กำหนด ────────────────────────────
# ถ้าเรียงผิด Word จะถือว่าไฟล์เสียแล้วซ่อมเอง ค่าที่ตั้งไว้จะหายเงียบ ๆ
RPR_ORDER = [
    'rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps', 'strike',
    'dstrike', 'outline', 'shadow', 'emboss', 'imprint', 'noProof', 'snapToGrid',
    'vanish', 'webHidden', 'color', 'spacing', 'w', 'kern', 'position', 'sz',
    'szCs', 'highlight', 'u', 'effect', 'bdr', 'shd', 'fitText', 'vertAlign',
    'rtl', 'cs', 'em', 'lang', 'eastAsianLayout', 'specVanish', 'oMath',
]

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,          # แบบละติน ยืดช่องว่างระหว่างคำ
    "thai_justify": WD_ALIGN_PARAGRAPH.THAI_JUSTIFY,  # ยืดตามกฎอักษรเชิงซ้อน
}

LINE = {
    "single": WD_LINE_SPACING.SINGLE,
    "one_half": WD_LINE_SPACING.ONE_POINT_FIVE,
    "double": WD_LINE_SPACING.DOUBLE,
}


# ── สเปกตั้งต้นสำหรับงานวิชาการไทย ────────────────────────────────────
# วัดจากเอกสารที่พิมพ์บน Word จริงและผ่านการตรวจของสถาบันมาแล้ว
# ถ้างานของคุณใช้ค่าอื่น ให้คัดลอกไปแก้ อย่าแก้ไฟล์นี้
SPEC_TH_ACADEMIC = {
    "font": "TH Sarabun New",
    "size": 16,
    "lang_bidi": "th-TH",
    "align": "thai_justify",
    "line_spacing": "single",
    "space_before": 0,
    "space_after": 0,
    "margins_cm": {"top": 3.81, "bottom": 2.54, "left": 3.81, "right": 2.54},
    "style_prefix": "ต.",
    # ชื่อสไตล์ -> ค่าที่ต่างจาก Normal
    #   size · bold · italic · color · align · first_indent · left_indent · keep_next
    "styles": {
        "หัวบท":            {"size": 18, "bold": True, "align": "center", "keep_next": True},
        "ชื่อบท":           {"size": 18, "bold": True, "align": "center", "keep_next": True},
        "หัวข้อ":           {"bold": True, "align": "left", "keep_next": True},
        "เนื้อความ":        {"first_indent": 1.52},
        "เนื้อความไม่เยื้อง": {},
        "ข้อ1":             {"first_indent": 1.52},
        "ข้อ2":             {"first_indent": 1.96},
        "ข้อ3":             {"first_indent": 2.54},
        "ข้อ4":             {"first_indent": 2.97},
        "หมายเหตุ":         {"size": 14, "italic": True, "color": "595959", "align": "left"},
        "ชื่อตาราง":        {"align": "left", "keep_next": True},
        "ข้อความในตาราง":   {"align": "center"},
        "ชื่อภาพ":          {"align": "center"},
    },
}


# ── ฟังก์ชันระดับ XML ────────────────────────────────────────────────

def _reorder_rpr(rpr):
    kids = list(rpr)
    kids.sort(key=lambda e: RPR_ORDER.index(e.tag.split('}')[1])
              if e.tag.split('}')[1] in RPR_ORDER else 999)
    for k in kids:
        rpr.append(k)


def _need(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def bind_complex_script(style, font, size_pt, lang_bidi, bold=False, italic=False):
    """ผูกฟอนต์ ขนาด ภาษา และธงอักษรเชิงซ้อน เข้ากับสไตล์

    ต้องตั้ง w:szCs คู่กับ w:sz และ w:bCs คู่กับ w:b เสมอ
    ไม่งั้นภาษาไทยจะไม่เปลี่ยนขนาดและไม่หนาตามที่สั่ง
    """
    rpr = style.element.get_or_add_rPr()

    rf = _need(rpr, 'w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(attr), font)

    if size_pt is not None:
        _need(rpr, 'w:szCs').set(qn('w:val'), str(int(size_pt * 2)))
    if bold:
        _need(rpr, 'w:bCs')
    if italic:
        _need(rpr, 'w:iCs')

    _need(rpr, 'w:cs')  # ธงอักษรเชิงซ้อน

    lang = _need(rpr, 'w:lang')
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:bidi'), lang_bidi)

    _reorder_rpr(rpr)


def set_document_language(doc, lang_bidi):
    """แก้ค่าภาษาระดับเอกสาร 3 จุด ที่ python-docx ตั้งผิดไว้"""
    for lang in doc.styles.element.findall('.//' + qn('w:lang')):
        lang.set(qn('w:bidi'), lang_bidi)

    settings = doc.settings.element

    tfl = settings.find(qn('w:themeFontLang'))
    if tfl is None:
        tfl = OxmlElement('w:themeFontLang')
        tfl.set(qn('w:val'), 'en-US')
        settings.append(tfl)
    tfl.set(qn('w:bidi'), lang_bidi)

    compat = settings.find(qn('w:compat'))
    if compat is None:
        compat = OxmlElement('w:compat')
        settings.insert(0, compat)
    if compat.find(qn('w:applyBreakingRules')) is None:
        compat.insert(0, OxmlElement('w:applyBreakingRules'))


def tag_runs_complex_script(doc, lang_bidi):
    """ติดธง <w:cs/> และภาษา ลงทุก run ในเอกสาร

    จุดนี้พลาดกันบ่อยที่สุด Word ไม่สืบทอดธงนี้จากสไตล์มาใช้ตัดบรรทัด
    ตั้งไว้แค่ที่สไตล์ไม่พอ ต้องเขียนลงระดับ run เหมือนที่ Word เขียนเอง
    """
    n = 0
    for r in doc.element.body.iter(qn('w:r')):
        rpr = r.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            r.insert(0, rpr)
        _need(rpr, 'w:cs')
        _need(rpr, 'w:lang').set(qn('w:bidi'), lang_bidi)
        _reorder_rpr(rpr)
        n += 1
    return n


# ── ตัวสร้างเอกสาร ───────────────────────────────────────────────────

class DocBuilder:
    """สร้างเอกสาร Word จากสเปกเดียว ทุกย่อหน้าผูกสไตล์ ไม่มีการจัดรูปแบบโดยตรง"""

    def __init__(self, spec=None, template=None):
        self.spec = copy.deepcopy(spec or SPEC_TH_ACADEMIC)
        self.doc = Document(template) if template else Document()
        self.prefix = self.spec.get("style_prefix", "")
        self._setup()

    # -- ตั้งค่า --------------------------------------------------------
    def _setup(self):
        sp = self.spec
        nm = self.doc.styles['Normal']
        nm.font.name = sp["font"]
        nm.font.size = Pt(sp["size"])
        pf = nm.paragraph_format
        pf.alignment = ALIGN[sp.get("align", "left")]
        ls = sp.get("line_spacing", "single")
        if isinstance(ls, str):
            pf.line_spacing_rule = LINE[ls]
        else:
            pf.line_spacing = ls
        pf.space_before = Pt(sp.get("space_before", 0))
        pf.space_after = Pt(sp.get("space_after", 0))
        bind_complex_script(nm, sp["font"], sp["size"], sp["lang_bidi"])

        set_document_language(self.doc, sp["lang_bidi"])

        for name, opt in sp.get("styles", {}).items():
            self._add_style(name, opt)

        m = sp.get("margins_cm")
        if m:
            sec = self.doc.sections[0]
            sec.top_margin = Cm(m["top"])
            sec.bottom_margin = Cm(m["bottom"])
            sec.left_margin = Cm(m["left"])
            sec.right_margin = Cm(m["right"])

    def _add_style(self, name, opt):
        sp = self.spec
        st = self.doc.styles.add_style(self.prefix + name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = self.doc.styles['Normal']
        st.quick_style = True          # ให้โผล่ในแถบสไตล์ของ Word
        size = opt.get("size", sp["size"])
        st.font.size = Pt(size)
        st.font.bold = opt.get("bold", False)
        st.font.italic = opt.get("italic", False)
        if opt.get("color"):
            st.font.color.rgb = RGBColor.from_string(opt["color"])
        pf = st.paragraph_format
        if opt.get("align"):
            pf.alignment = ALIGN[opt["align"]]
        if opt.get("first_indent") is not None:
            pf.first_line_indent = Cm(opt["first_indent"])
        if opt.get("left_indent") is not None:
            pf.left_indent = Cm(opt["left_indent"])
        pf.keep_with_next = opt.get("keep_next", False)
        bind_complex_script(st, sp["font"], size, sp["lang_bidi"],
                            opt.get("bold", False), opt.get("italic", False))
        return st

    # -- เขียนเนื้อหา ---------------------------------------------------
    def p(self, text="", style="เนื้อความ"):
        """ย่อหน้าหนึ่งย่อหน้า ผูกสไตล์อย่างเดียว"""
        return self.doc.add_paragraph(text, style=self.prefix + style)

    def body(self, text=""):
        return self.p(text, "เนื้อความ")

    def note(self, text):
        """บรรทัดคำแนะนำสีเทา สำหรับร่างระหว่างทาง ลบทิ้งตอนจัดเล่มจริง"""
        return self.p(text, "หมายเหตุ")

    def blank(self):
        return self.p("", "เนื้อความไม่เยื้อง")

    def page_break(self):
        par = self.doc.add_paragraph(style=self.prefix + "เนื้อความไม่เยื้อง")
        par.add_run().add_break(WD_BREAK.PAGE)
        return par

    def table(self, caption, headers, rows, style="Table Grid"):
        """ตารางมาตรฐาน — คำบรรยายเหนือตาราง หัวซ้ำข้ามหน้า พอดีหน้ากระดาษ

        คำบรรยาย**ตาราง**อยู่เหนือ ส่วนคำบรรยาย**ภาพ**อยู่ใต้ ตามหลัก APA 7th
        """
        if caption:
            self.p(caption, "ชื่อตาราง")

        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = style

        tblPr = t._tbl.tblPr
        for tag in ('w:tblW', 'w:tblLayout'):
            old = tblPr.find(qn(tag))
            if old is not None:
                tblPr.remove(old)
        w = OxmlElement('w:tblW')
        w.set(qn('w:type'), 'pct')
        w.set(qn('w:w'), '5000')          # 5000 = 100% = AutoFit Window
        tblPr.append(w)
        lay = OxmlElement('w:tblLayout')
        lay.set(qn('w:type'), 'autofit')
        tblPr.append(lay)

        def fill(cell, text, bold=False):
            par = cell.paragraphs[0]
            par.style = self.doc.styles[self.prefix + "ข้อความในตาราง"]
            par.add_run(str(text)).bold = bold

        for c, h in enumerate(headers):
            fill(t.rows[0].cells[c], h, bold=True)

        # หัวตารางซ้ำเมื่อตารางข้ามหน้า
        trPr = t.rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            trPr.append(OxmlElement('w:tblHeader'))

        for row in rows:
            cells = t.add_row().cells
            for c, v in enumerate(row):
                fill(cells[c], v)
        return t

    def figure_placeholder(self, number, title, layout, labels="", center="",
                           source="", note=""):
        """ตัวยึดตำแหน่งภาพ พร้อมบล็อกสเปกให้คัดลอกไปสั่งผลิตภาพต่อ

        วางตรงจุดที่เนื้อความพูดถึงภาพนั้น ห้ามกองรวมไว้ท้ายหัวข้อ
        """
        self.blank()
        self.p("[ภาพประกอบ %s]" % number, "ชื่อภาพ")
        lines = ["ภาพที่ %s %s" % (number, title), "รูปแบบ: %s" % layout]
        if labels or center:
            lines.append("ป้ายกำกับ: %s | ศูนย์กลาง: %s" % (labels, center))
        if source:
            lines.append("ที่มาแนวคิด: %s" % source)
        if note:
            lines.append("หมายเหตุ: %s" % note)
        for ln in lines:
            self.p(ln, "หมายเหตุ")
        self.blank()

    # -- บันทึก ---------------------------------------------------------
    def save(self, path, overwrite=True):
        """บันทึกไฟล์ พร้อมติดธงอักษรเชิงซ้อนลงทุก run ก่อนเสมอ"""
        tag_runs_complex_script(self.doc, self.spec["lang_bidi"])
        if not overwrite and os.path.exists(path):
            raise FileExistsError(path)
        try:
            self.doc.save(path)
        except PermissionError:
            raise PermissionError(
                "เขียนทับไม่ได้ ไฟล์เปิดค้างอยู่ในโปรแกรมอื่น ให้ปิดก่อน: %s" % path)
        return path


# ── ด่านตรวจ ─────────────────────────────────────────────────────────

def verify(path, lang_bidi="th-TH"):
    """ตรวจไฟล์ที่ผลิตแล้ว คืนค่า (ผ่านหรือไม่, รายการผลตรวจ)"""
    import zipfile
    z = zipfile.ZipFile(path)
    d = z.read("word/document.xml").decode("utf-8")
    st = z.read("word/styles.xml").decode("utf-8")
    se = z.read("word/settings.xml").decode("utf-8")
    ntbl = d.count("<w:tbl>")

    items = [
        ("ธงอักษรเชิงซ้อนครบทุก run",
         d.count("<w:cs/>") == d.count("<w:r>") and d.count("<w:r>") > 0,
         "%d/%d" % (d.count("<w:cs/>"), d.count("<w:r>"))),
        ("ไม่เหลือภาษาเริ่มต้น ar-SA", st.count("ar-SA") == 0, str(st.count("ar-SA"))),
        ("ตั้งภาษาอักษรเชิงซ้อนถูกต้อง", lang_bidi in st, lang_bidi),
        ("เปิดกฎการตัดบรรทัด", "applyBreakingRules" in se, "-"),
        ("หัวตารางซ้ำเมื่อข้ามหน้า",
         ntbl == 0 or d.count("<w:tblHeader/>") >= ntbl,
         "%d/%d" % (d.count("<w:tblHeader/>"), ntbl)),
        ("ตารางพอดีหน้ากระดาษ",
         ntbl == 0 or d.count('w:w="5000"') >= ntbl,
         "%d/%d" % (d.count('w:w="5000"'), ntbl)),
        ("ตารางปรับความกว้างอัตโนมัติ",
         ntbl == 0 or d.count('w:tblLayout w:type="autofit"') >= ntbl,
         "%d/%d" % (d.count('w:tblLayout w:type="autofit"'), ntbl)),
        ("ไม่ซ้อนเลขข้อ",
         not any("%d. %d." % (i, i) in d for i in range(1, 10)), "-"),
        ("ไม่จำลองการเยื้องด้วย tab", d.count("<w:tab/>") == 0,
         str(d.count("<w:tab/>"))),
    ]
    return all(ok for _, ok, _ in items), items


def measure(path, limit=60):
    """วัดสเปกจากเอกสารตัวอย่างที่ผ่านการยอมรับแล้ว

    ใช้ก่อนเริ่มงานใหม่ทุกครั้ง แม่นกว่าการตีความคู่มือ
    ระยะบรรทัดเป็น None แปลว่าเดี่ยว · การจัดวางเป็น None แปลว่าชิดซ้าย
    """
    d = Document(path)
    sec = d.sections[0]
    out = ["ขอบกระดาษ ซม. บน %.2f ล่าง %.2f ซ้าย %.2f ขวา %.2f" % (
        sec.top_margin.cm, sec.bottom_margin.cm, sec.left_margin.cm, sec.right_margin.cm)]
    shown = 0
    for i, par in enumerate(d.paragraphs):
        if not par.text.strip():
            continue
        if shown >= limit:
            break
        shown += 1
        pf = par.paragraph_format
        r = par.runs[0] if par.runs else None
        out.append("%3d | %-14s บรรทัดแรก %-6s ซ้าย %-6s บรรทัด %-5s จัดวาง %-8s "
                   "ขนาด %-5s หนา %-5s | %s" % (
            i, par.style.name,
            round(pf.first_line_indent.cm, 2) if pf.first_line_indent is not None else "-",
            round(pf.left_indent.cm, 2) if pf.left_indent is not None else "-",
            pf.line_spacing,
            str(pf.alignment).split(" ")[0] if pf.alignment else "-",
            r.font.size.pt if r and r.font.size else None,
            r.bold if r else None, par.text[:36]))
    return "\n".join(out)
