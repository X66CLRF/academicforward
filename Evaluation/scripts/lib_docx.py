# -*- coding: utf-8 -*-
"""Shared helpers for the charnankan-doc skill.

Thai literals are written as \\u escapes on purpose: the scripts must behave
identically no matter what console codepage Windows hands us.
"""
import copy
import datetime
import os
import re
import shutil

import docx

# "เอกสารแนบที่"
ATTACH = "เอกสารแนบที่"
# "(ดูเอกสารแนบที่ "
CITE_OPEN = "(ดู" + ATTACH + " "
THAI_DIGITS = "๐๑๒๓๔๕๖๗๘๙"

DIGITS = "0-9" + THAI_DIGITS
NUM = "[" + DIGITS + "]+"
CAPTION_RE = re.compile(ATTACH + r"\s*(" + NUM + ")")
CITE_RE = re.compile(re.escape(CITE_OPEN.rstrip()) + r"\s*([" + DIGITS + r",\s\-]+)\)")

# The house style already used in the document:
#   "เอกสารแนบ<หมวด> ระดับที่ ๓: หมายเลข ๔๔, ๔๕ และ ๗๑"
HOUSE_RE = re.compile("หมายเลข" + r"\s*((?:" + NUM + r"[,\s]*|และ\s*)+)")
# The prefix in front of that list, captured whole so it can be reused verbatim.
PREFIX_RE = re.compile("(" + ATTACH[:-3] + r"[^:\n]*?ระดับที่\s*" + NUM + r")\s*:\s*หมายเลข")
LEVEL_RE = re.compile("ระดับที่\\s*(" + NUM + ")")


def to_int(s):
    """Read a number written with Arabic or Thai digits."""
    out = ""
    for ch in s:
        i = THAI_DIGITS.find(ch)
        out += str(i) if i >= 0 else ch
    return int(out)


def load(path):
    return docx.Document(path)


def backup(path):
    """Copy the file next to itself before any write. Never skipped."""
    stamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    root, ext = os.path.splitext(path)
    dest = "%s.bak-%s%s" % (root, stamp, ext)
    shutil.copy2(path, dest)
    return dest


def captions(doc):
    """Every attachment caption, in document order.

    Returns [(index_in_body, paragraph, number)].
    """
    out = []
    for i, p in enumerate(doc.paragraphs):
        m = CAPTION_RE.match(p.text.strip())
        if m:
            out.append((i, p, to_int(m.group(1))))
    return out


def criteria_cells(doc):
    """Column-3 cells ('บันทึกร่องรอยคุณภาพ') of every 3-column table.

    Returns [(table_index, row_index, cell)] for cells that hold text.
    """
    out = []
    for ti, t in enumerate(doc.tables):
        if len(t.columns) != 3:
            continue
        for ri, row in enumerate(t.rows):
            cell = row.cells[2]
            if cell.text.strip():
                out.append((ti, ri, cell))
    return out


def cell_runs(cell):
    return [r for p in cell.paragraphs for r in p.runs]


def plain_run(cell):
    """A non-bold run to clone formatting from; falls back to any run."""
    runs = [r for r in cell_runs(cell) if r.text]
    for r in runs:
        if not r.bold:
            return r
    return runs[-1] if runs else None


def set_paragraph_text(par, text):
    """Replace a paragraph's text while keeping its formatting exactly.

    Writes into the first run that carries text and blanks the rest. Safe only
    when every run in the paragraph shares one format - check with
    is_uniform() first.
    """
    runs = par.runs
    target = None
    for r in runs:
        if r.text:
            target = r
            break
    if target is None:
        if not runs:
            par.add_run(text)
            return
        target = runs[0]
    target.text = text
    for r in runs:
        if r is not target:
            r.text = ""


def run_sig(r):
    f = r.font
    color = None
    if f.color is not None and f.color.type is not None:
        color = str(f.color.rgb)
    return (r.bold, r.italic, r.underline, f.size, f.name, color)


def is_uniform(par):
    sigs = {run_sig(r) for r in par.runs if r.text}
    return len(sigs) <= 1


def read_citation(cell):
    """Numbers already cited in this cell, or None if it carries no citation.

    Reads both the house style ("... หมายเลข ๔๔, ๔๕ และ ๗๑") and the short
    style ("(ดูเอกสารแนบที่ 44, 45, 71)"). A cell may hold several house-style
    lists; every number found is returned.
    """
    found = []
    text = cell.text
    m = CITE_RE.search(text)
    if m:
        found += re.findall(NUM, m.group(1))
    for group in HOUSE_RE.findall(text):
        found += re.findall(NUM, group)
    if not found:
        return None
    seen = []
    for x in found:
        n = to_int(x)
        if n not in seen:
            seen.append(n)
    return seen


def read_prefix(cell):
    """The house-style prefix this cell already uses, if any.

    e.g. "เอกสารแนบการใช้คอมพิวเตอร์ ระดับที่ ๓"
    """
    m = PREFIX_RE.search(" ".join(cell.text.split()))
    return m.group(1) if m else None


def read_level(row):
    """The level written in column 2 of a criteria row, as Thai digits."""
    m = LEVEL_RE.search(" ".join(row.cells[1].text.split()))
    return m.group(1) if m else None


def format_citation(prefix, numbers):
    """Render the house style: '<prefix>: หมายเลข ๔๔, ๔๕ และ ๗๑'."""
    nums = [str(n) for n in numbers]
    if len(nums) == 1:
        body = nums[0]
    else:
        body = ", ".join(nums[:-1]) + " และ " + nums[-1]
    return prefix + ": " + "หมายเลข" + " " + body


def write_citation(cell, numbers, prefix=None):
    """Append a citation to the end of a cell.

    Appends a run cloned from an existing plain run, so the citation inherits
    the cell's font instead of Word's defaults. With no prefix it falls back to
    the short style, which is only used when the house prefix is unknown.
    """
    if prefix:
        text = format_citation(prefix, numbers)
    else:
        text = CITE_OPEN + ", ".join(str(n) for n in numbers) + ")"

    # Drop an existing short-style citation first, wherever it sits.
    for p in cell.paragraphs:
        for r in p.runs:
            if CITE_OPEN.rstrip() in r.text or CITE_RE.search(r.text):
                r.text = CITE_RE.sub("", r.text).rstrip()

    par = cell.paragraphs[-1]
    src = plain_run(cell)
    if src is None:
        par.add_run(" " + text)
        return
    new = copy.deepcopy(src._r)
    par._p.append(new)
    from docx.text.run import Run
    Run(new, par).text = " " + text
